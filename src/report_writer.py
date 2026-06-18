from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

from .advanced_diagnostics import build_advanced_diagnostics
from .analysis_engine import (
    build_dupont_driver_table,
    build_management_issue_table,
    build_profit_bridge_table,
    build_sensitivity_risk_table,
)
from .assignment_filters import check_assignment_conditions
from .charts import create_charts
from .company_master import select_companies
from .config_loader import PROJECT_ROOT, load_industry_policy, load_rubric
from .course_framework import (
    build_plus_alpha_analysis_table,
    build_plus_alpha_commentary,
    build_required_plus_alpha_table,
)
from .data_loader import Dataset, load_dataset
from .metrics.financial import KEY_METRICS, METRIC_LABELS, compute_financial_metrics, latest_metrics
from .metrics.scoring import SCORE_LABELS, build_company_scores, build_scoring_notes
from .narrative import (
    build_alpha_commentary,
    build_cashflow_commentary,
    build_profitability_commentary,
    build_selection_reason,
    build_stability_commentary,
    business_descriptions,
    causal_matrix,
    collect_missing_notes,
    nine_perspectives,
)
from .rubric_checker import build_assignment_response_table, required_sections_table, sanitize_advice_terms


PERCENT_COLUMNS = {
    "revenue_growth_rate",
    "operating_margin",
    "roa",
    "roe",
    "roa_decomposed",
    "roe_decomposed",
    "net_margin",
    "current_ratio",
    "equity_ratio",
    "debt_ratio",
    "fixed_ratio",
    "fixed_long_term_adequacy_ratio",
    "safety_margin",
}

NUMBER_COLUMNS = {
    "revenue",
    "operating_income",
    "net_income",
    "total_assets",
    "equity",
    "cash_flow_operating",
    "fcf",
    "asset_turnover",
    "financial_leverage",
    "break_even_sales",
    "operating_leverage",
    "per",
    "pbr",
}

SCORE_COLUMNS = {
    "data_completeness",
    "growth_score",
    "profitability_score",
    "stability_score",
    "cashflow_score",
    "efficiency_score",
    "trend_resilience_score",
    "valuation_reference_score",
    "analysis_quality_score",
}


@dataclass
class ReportPackage:
    docx_path: Path
    chart_paths: dict[str, Path]
    warnings: list[str]
    notes: list[str]
    missing_notes: list[str]
    condition_table: pd.DataFrame
    company_check_table: pd.DataFrame
    assignment_response_table: pd.DataFrame
    metrics: pd.DataFrame
    quality_scores: pd.DataFrame


def _missing_label(rubric: dict[str, Any]) -> str:
    return str(rubric["assignment"]["missing_value_label"])


def _format_value(value: object, column: str, missing_label: str) -> str:
    if value is None or pd.isna(value):
        return missing_label
    if column in PERCENT_COLUMNS:
        numeric = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
        return missing_label if pd.isna(numeric) else f"{numeric * 100:.1f}%"
    if column in NUMBER_COLUMNS:
        numeric = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
        if pd.isna(numeric):
            return missing_label
        if abs(float(numeric)) >= 100:
            return f"{numeric:,.0f}"
        return f"{numeric:,.2f}"
    if column in SCORE_COLUMNS:
        numeric = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
        return missing_label if pd.isna(numeric) else f"{numeric:.1f}"
    return str(value)


def _display_table(
    df: pd.DataFrame,
    columns: list[str],
    labels: dict[str, str],
    missing_label: str,
) -> pd.DataFrame:
    rows: list[dict[str, str]] = []
    for _, row in df.iterrows():
        display_row = {}
        for column in columns:
            display_row[labels.get(column, column)] = _format_value(row.get(column), column, missing_label)
        rows.append(display_row)
    return pd.DataFrame(rows)


def _add_df_table(document: Document, df: pd.DataFrame, missing_label: str) -> None:
    if df.empty:
        document.add_paragraph(missing_label)
        return
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(df.columns):
        table.rows[0].cells[idx].text = str(column)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(df.columns):
            value = row[column]
            cells[idx].text = missing_label if value is None or pd.isna(value) else str(value)


def _add_paragraph(document: Document, text: str, rubric: dict[str, Any]) -> None:
    document.add_paragraph(sanitize_advice_terms(text, rubric))


def _set_document_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)


def _mode_label(app_mode: str) -> str:
    return "課題モード" if app_mode == "assignment" else "汎用モード"


def _latest_with_company(metrics: pd.DataFrame) -> pd.DataFrame:
    latest = latest_metrics(metrics)
    return latest.sort_values(["_selection_order", "ticker"]).reset_index(drop=True)


def _write_report(
    docx_path: Path,
    preset: dict[str, Any],
    app_mode: str,
    industry_mode: str,
    companies: pd.DataFrame,
    metrics: pd.DataFrame,
    chart_paths: dict[str, Path],
    assignment_result: dict[str, object],
    rubric: dict[str, Any],
    as_of: date,
) -> None:
    missing_label = _missing_label(rubric)
    assignment_table = build_assignment_response_table(rubric)
    latest = _latest_with_company(metrics)
    quality_scores = build_company_scores(metrics)
    advanced = build_advanced_diagnostics(
        metrics,
        companies,
        app_mode=app_mode,
        industry_mode=industry_mode,
    )
    alpha_table = build_alpha_commentary(latest, missing_label)
    framework_table = build_required_plus_alpha_table()
    plus_alpha_table = build_plus_alpha_analysis_table(metrics, companies, missing_label=missing_label)
    dupont_table = build_dupont_driver_table(metrics, missing_label)
    profit_bridge_table = build_profit_bridge_table(metrics, missing_label)
    sensitivity_table = build_sensitivity_risk_table(metrics, missing_label=missing_label)
    management_issue_table = build_management_issue_table(
        metrics,
        companies,
        app_mode=app_mode,
        industry_mode=industry_mode,
        missing_label=missing_label,
    )
    missing_notes = collect_missing_notes(metrics, missing_label)

    document = Document()
    _set_document_style(document)

    document.add_heading("日本上場企業比較レポート", level=0)
    _add_paragraph(document, f"プリセット: {preset.get('name', preset.get('preset_id', ''))}", rubric)
    _add_paragraph(document, f"分析モード: {_mode_label(app_mode)}", rubric)
    _add_paragraph(document, f"業種判定モード: {industry_mode}", rubric)
    _add_paragraph(document, f"作成日: {as_of.isoformat()}", rubric)
    _add_paragraph(document, "本レポートは学習目的の比較分析であり、株式取引の推奨を目的としない。", rubric)

    document.add_heading("課題対応表", level=1)
    _add_df_table(document, assignment_table, missing_label)

    document.add_heading("条件適合表", level=1)
    _add_df_table(document, assignment_result["condition_table"], missing_label)
    document.add_paragraph("企業別判定")
    _add_df_table(document, assignment_result["company_check_table"], missing_label)

    warnings = assignment_result["warnings"]
    if warnings:
        document.add_paragraph("警告")
        for warning in warnings:
            _add_paragraph(document, f"・{warning}", rubric)

    notes = assignment_result["notes"]
    if notes:
        document.add_paragraph("注記")
        for note in notes:
            _add_paragraph(document, f"・{note}", rubric)

    document.add_heading("企業選定理由", level=1)
    _add_paragraph(
        document,
        build_selection_reason(companies, preset, assignment_result["industry_result"]),
        rubric,
    )

    document.add_heading("事業内容", level=1)
    _add_df_table(document, business_descriptions(companies), missing_label)

    document.add_heading("講義フレームワーク", level=1)
    document.add_heading("必須部分と＋αの区分", level=2)
    _add_df_table(document, framework_table, missing_label)
    document.add_heading("因果のマトリクス", level=2)
    _add_df_table(document, causal_matrix(companies), missing_label)
    document.add_heading("企業分析9視点", level=2)
    _add_df_table(document, nine_perspectives(companies), missing_label)

    document.add_heading("主要財務数値表", level=1)
    financial_columns = [
        "ticker",
        "company_name",
        "fiscal_year",
        "revenue",
        "operating_income",
        "net_income",
        "total_assets",
        "equity",
        "cash_flow_operating",
        "fcf",
    ]
    financial_labels = {
        "ticker": "証券コード",
        "company_name": "企業名",
        "fiscal_year": "年度",
        "revenue": "売上高",
        "operating_income": "営業利益",
        "net_income": "当期利益",
        "total_assets": "総資産",
        "equity": "自己資本",
        "cash_flow_operating": "営業CF",
        "fcf": "FCF",
    }
    _add_df_table(document, _display_table(metrics, financial_columns, financial_labels, missing_label), missing_label)

    document.add_heading("財務指標表", level=1)
    metric_columns = [
        "ticker",
        "company_name",
        "fiscal_year",
        *KEY_METRICS,
        "roa_decomposed",
        "roe_decomposed",
    ]
    metric_labels = {
        "ticker": "証券コード",
        "company_name": "企業名",
        "fiscal_year": "年度",
        **METRIC_LABELS,
        "roa_decomposed": "ROA分解",
        "roe_decomposed": "ROE分解",
    }
    _add_df_table(document, _display_table(metrics, metric_columns, metric_labels, missing_label), missing_label)

    document.add_heading("高度分析サマリー", level=1)
    score_columns = [
        "ticker",
        "company_name",
        "fiscal_year",
        "data_completeness",
        "growth_score",
        "profitability_score",
        "stability_score",
        "cashflow_score",
        "efficiency_score",
        "trend_resilience_score",
        "valuation_reference_score",
        "analysis_quality_score",
        "analysis_band",
    ]
    score_labels = {
        "ticker": "証券コード",
        "company_name": "企業名",
        "fiscal_year": "年度",
        **SCORE_LABELS,
    }
    score_columns = [column for column in score_columns if column in quality_scores.columns]
    _add_df_table(document, _display_table(quality_scores, score_columns, score_labels, missing_label), missing_label)
    for note in build_scoring_notes(quality_scores, missing_label):
        _add_paragraph(document, note, rubric)

    document.add_heading("高度判定・考察", level=1)
    _add_df_table(document, advanced["diagnostic_table"], missing_label)
    for paragraph in advanced["commentary"]:
        _add_paragraph(document, str(paragraph), rubric)
    if advanced["mode_notes"]:
        document.add_paragraph("モード別確認観点")
        for note in advanced["mode_notes"]:
            _add_paragraph(document, f"・{note}", rubric)

    document.add_heading("高度アルゴリズム分析", level=1)
    document.add_paragraph("ROE要因分解")
    _add_df_table(document, dupont_table, missing_label)
    document.add_paragraph("営業利益ブリッジ")
    _add_df_table(document, profit_bridge_table, missing_label)
    document.add_paragraph("感応度・リスクフラグ")
    _add_df_table(document, sensitivity_table, missing_label)
    document.add_paragraph("経営分析論点")
    _add_df_table(document, management_issue_table, missing_label)

    document.add_heading("グラフ", level=1)
    chart_titles = {
        "revenue_trend": "売上高推移",
        "operating_margin_trend": "営業利益率推移",
        "roa_roe_trend": "ROA/ROE推移",
        "equity_ratio_trend": "自己資本比率推移",
        "cashflow_fcf_trend": "営業CF/FCF推移",
    }
    for slug, path in chart_paths.items():
        document.add_paragraph(chart_titles.get(slug, slug))
        if path.exists():
            document.add_picture(str(path), width=Inches(6.3))
        else:
            document.add_paragraph(missing_label)

    document.add_heading("収益性分析", level=1)
    for paragraph in build_profitability_commentary(latest, missing_label):
        _add_paragraph(document, paragraph, rubric)

    document.add_heading("財務安定性分析", level=1)
    for paragraph in build_stability_commentary(latest, missing_label):
        _add_paragraph(document, paragraph, rubric)

    document.add_heading("キャッシュフロー分析", level=1)
    for paragraph in build_cashflow_commentary(latest, missing_label):
        _add_paragraph(document, paragraph, rubric)

    document.add_heading("＋α分析", level=1)
    for paragraph in build_plus_alpha_commentary(plus_alpha_table, missing_label):
        _add_paragraph(document, paragraph, rubric)
    _add_df_table(document, alpha_table, missing_label)
    document.add_paragraph("講義手法に基づく＋α詳細")
    _add_df_table(document, plus_alpha_table, missing_label)

    document.add_heading("総合比較", level=1)
    _add_paragraph(
        document,
        "各社の差は、収益性、資産効率、財務安定性、キャッシュ創出力、事業モデルの組み合わせとして整理できる。"
        "本MVPではサンプルCSVに基づき、同じ表形式で他業種にも展開できるようにした。",
        rubric,
    )

    document.add_heading("結論", level=1)
    _add_paragraph(
        document,
        "課題条件の適合状況と財務指標を分けて確認することで、比較対象の妥当性と分析結果を混同せずに扱える。"
        "最終提出では、サンプルデータを一次資料に差し替え、欠損箇所の根拠を確認する必要がある。",
        rubric,
    )

    document.add_heading("参考資料", level=1)
    references = [
        "data/company_master/sample_company_master.csv",
        "data/sample_financials/sample_financials.csv",
        "data/sample_financials/sample_market_data.csv",
        "data/manual_kpis/sample_manual_kpis.csv",
        "config/rubric/assignment_rubric.yaml",
        "config/industry_policy.yaml",
        "EDINET API雛形: src/edinet_client.py",
    ]
    for reference in references:
        _add_paragraph(document, f"・{reference}", rubric)

    document.add_heading("欠損データ注記", level=1)
    if missing_notes:
        for note in missing_notes:
            _add_paragraph(document, f"・{note}", rubric)
    else:
        _add_paragraph(document, "主要指標の最新年度について、MVPサンプル上の欠損注記はありません。", rubric)

    document.add_heading("出力章チェック", level=1)
    _add_df_table(document, required_sections_table(rubric), missing_label)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def build_report_package(
    selected_tickers: list[str],
    preset: dict[str, Any],
    app_mode: str = "assignment",
    industry_mode: str | None = None,
    dataset: Dataset | None = None,
    output_dir: Path = PROJECT_ROOT / "output",
    as_of: date | None = None,
) -> ReportPackage:
    rubric = load_rubric()
    industry_policy = load_industry_policy()
    dataset = dataset or load_dataset(use_sqlite=True)
    selected_companies = select_companies(dataset.company_master, selected_tickers).copy()
    selected_companies["_selection_order"] = range(len(selected_companies))
    industry_mode = industry_mode or str(preset.get("industry_mode") or rubric["assignment"]["default_industry_mode"])
    as_of = as_of or date.today()

    assignment_result = check_assignment_conditions(
        selected_companies,
        app_mode=app_mode,
        industry_mode=industry_mode,
        rubric=rubric,
        industry_policy=industry_policy,
        as_of=as_of,
    )

    selected_financials = dataset.financials[dataset.financials["ticker"].isin(selected_companies["ticker"])].copy()
    selected_market = dataset.market_data[dataset.market_data["ticker"].isin(selected_companies["ticker"])].copy()
    selected_manual = dataset.manual_kpis[dataset.manual_kpis["ticker"].isin(selected_companies["ticker"])].copy()
    metrics = compute_financial_metrics(selected_financials, selected_market, selected_manual)
    metrics = metrics.merge(
        selected_companies[["ticker", "company_name", "_selection_order"]],
        on="ticker",
        how="left",
    )
    metrics = metrics.sort_values(["_selection_order", "fiscal_year"]).reset_index(drop=True)
    quality_scores = build_company_scores(metrics)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    preset_id = str(preset.get("preset_id", "custom"))
    chart_dir = output_dir / "charts" / f"{preset_id}_{timestamp}"
    chart_paths = create_charts(metrics, selected_companies, chart_dir)
    docx_path = output_dir / "reports" / f"{preset_id}_{timestamp}.docx"

    _write_report(
        docx_path=docx_path,
        preset=preset,
        app_mode=app_mode,
        industry_mode=industry_mode,
        companies=selected_companies,
        metrics=metrics,
        chart_paths=chart_paths,
        assignment_result=assignment_result,
        rubric=rubric,
        as_of=as_of,
    )

    missing_notes = collect_missing_notes(metrics, _missing_label(rubric))
    return ReportPackage(
        docx_path=docx_path,
        chart_paths=chart_paths,
        warnings=list(assignment_result["warnings"]),
        notes=list(assignment_result["notes"]),
        missing_notes=missing_notes,
        condition_table=assignment_result["condition_table"],
        company_check_table=assignment_result["company_check_table"],
        assignment_response_table=build_assignment_response_table(rubric),
        metrics=metrics,
        quality_scores=quality_scores,
    )
