from datetime import date
from pathlib import Path
from uuid import uuid4

import pandas as pd
from docx import Document

from src.config_loader import PROJECT_ROOT
from src.config_loader import load_presets
from src.data_loader import Dataset, load_sample_dataset
from src.report_writer import REPORT_TYPE_DIAGNOSTIC, REPORT_TYPE_SUBMISSION, build_report_package


def _document_text(path):
    document = Document(path)
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text)
    return "\n".join(paragraph_text + table_text)


def _heading1_texts(path):
    document = Document(path)
    return [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 1"]


def _header_text(path):
    document = Document(path)
    parts = []
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.header.paragraphs)
    return "\n".join(parts)


def _footer_text(path):
    document = Document(path)
    parts = []
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(parts)


def _footer_xml(path):
    document = Document(path)
    return "\n".join(section.footer._element.xml for section in document.sections)


def test_report_contains_required_sections_and_warning():
    preset = load_presets()["friend_cafe_theme"]
    output_dir = PROJECT_ROOT / "work" / "pytest_reports" / uuid4().hex
    package = build_report_package(
        selected_tickers=preset["companies"],
        preset={**preset, "preset_id": "friend_cafe_theme"},
        app_mode="assignment",
        industry_mode=preset["industry_mode"],
        dataset=load_sample_dataset(),
        output_dir=output_dir,
        as_of=date(2026, 6, 17),
        edinet_filings=pd.DataFrame(
            [
                {
                    "doc_id": "S100TEST",
                    "edinet_code": "E32815",
                    "sec_code": "35430",
                    "filer_name": "テスト提出者",
                    "doc_description": "有価証券報告書",
                    "submit_datetime": "2026-06-18 10:00",
                    "csv_flag": "1",
                }
            ]
        ),
    )

    assert package.docx_path.exists()
    assert all(path.exists() for path in package.chart_paths.values())
    document = Document(package.docx_path)
    assert document.tables
    assert len(document.inline_shapes) >= 1
    text = _document_text(package.docx_path)
    for heading in [
        "課題対応表",
        "必須部分と＋αの区分",
        "＋α分析",
        "高度アルゴリズム分析",
        "経営分析論点",
        "警告",
        "参考資料",
        "欠損データ注記",
    ]:
        assert heading in text
    assert "JPX業種" in text
    assert "S100TEST" in text
    assert "E32815" in text
    assert "投資すべき" not in text


def test_submission_candidate_report_contains_only_final_selected_companies():
    preset = load_presets()["strict_cafe_retail"]
    output_dir = PROJECT_ROOT / "work" / "pytest_reports" / uuid4().hex
    package = build_report_package(
        selected_tickers=["3087", "3395"],
        preset={**preset, "preset_id": "strict_cafe_retail"},
        app_mode="assignment",
        industry_mode=preset["industry_mode"],
        dataset=load_sample_dataset(),
        output_dir=output_dir,
        as_of=date(2026, 6, 19),
        edinet_filings=pd.DataFrame(
            [
                {
                    "doc_id": "S100KOMEDA",
                    "edinet_code": "E32815",
                    "sec_code": "35430",
                    "filer_name": "コメダホールディングス",
                    "doc_description": "有価証券報告書",
                    "submit_datetime": "2026-06-18 10:00",
                    "csv_flag": "1",
                }
            ]
        ),
        report_type=REPORT_TYPE_SUBMISSION,
    )

    document = Document(package.docx_path)
    text = _document_text(package.docx_path)
    headings = _heading1_texts(package.docx_path)
    footer_xml = _footer_xml(package.docx_path)

    assert "3087" in text
    assert "3395" in text
    assert "コメダ" not in text
    assert "コメダホールディングス" not in text
    assert "3543" not in text
    for forbidden in ["生成前チェック", "使用データ表", "ツール", "アプリ", "edinet_candidate", "sample"]:
        assert forbidden not in text
    assert len(headings) == 14
    assert not any(heading.startswith("15.") for heading in headings)
    assert "14. 参考文献・出典" in headings
    reference_index = text.index("14. 参考文献・出典")
    assert "3543" not in text[reference_index:]
    assert "コメダ" not in text[reference_index:]
    assert "同業種上場企業における経営状況の比較分析" not in _header_text(package.docx_path)
    assert _footer_text(package.docx_path).strip() == ""
    assert "PAGE" in footer_xml
    assert len(document.tables) >= 5
    assert len(document.inline_shapes) >= 2


def test_diagnostic_report_can_show_pre_generation_checks():
    preset = load_presets()["strict_cafe_retail"]
    output_dir = PROJECT_ROOT / "work" / "pytest_reports" / uuid4().hex
    package = build_report_package(
        selected_tickers=["3087", "3395"],
        preset={**preset, "preset_id": "strict_cafe_retail"},
        app_mode="assignment",
        industry_mode=preset["industry_mode"],
        dataset=load_sample_dataset(),
        output_dir=output_dir,
        as_of=date(2026, 6, 19),
        report_type=REPORT_TYPE_DIAGNOSTIC,
    )

    text = _document_text(package.docx_path)
    assert "生成前チェック結果・使用データ表" in text


def _sparse_manual_dataset() -> Dataset:
    company_master = pd.DataFrame(
        [
            {
                "ticker": "1001",
                "company_name": "テスト自動車",
                "edinet_code": "E10001",
                "jpx_industry": "輸送用機器",
                "broad_sector": "自動車",
                "business_theme": "自動車",
                "listing_date": "1999-01-01",
                "listing_note": "",
                "business_summary": "自動車関連のテスト企業。",
                "source_note": "test",
            },
            {
                "ticker": "1002",
                "company_name": "テスト工業",
                "edinet_code": "E10002",
                "jpx_industry": "機械",
                "broad_sector": "機械",
                "business_theme": "重工",
                "listing_date": "2025-01-01",
                "listing_note": "",
                "business_summary": "機械関連のテスト企業。",
                "source_note": "test",
            },
        ]
    )
    financials = pd.DataFrame(
        [
            {
                "ticker": "1001",
                "fiscal_year": 2024,
                "revenue": 1000,
                "operating_income": pd.NA,
                "net_income": pd.NA,
                "total_assets": pd.NA,
                "equity": pd.NA,
                "current_assets": pd.NA,
                "current_liabilities": pd.NA,
                "fixed_assets": pd.NA,
                "long_term_liabilities": pd.NA,
                "cash_flow_operating": 80,
                "capex": pd.NA,
                "shares_outstanding": pd.NA,
            },
            {
                "ticker": "1002",
                "fiscal_year": 2024,
                "revenue": pd.NA,
                "operating_income": pd.NA,
                "net_income": pd.NA,
                "total_assets": pd.NA,
                "equity": pd.NA,
                "current_assets": pd.NA,
                "current_liabilities": pd.NA,
                "fixed_assets": pd.NA,
                "long_term_liabilities": pd.NA,
                "cash_flow_operating": pd.NA,
                "capex": pd.NA,
                "shares_outstanding": pd.NA,
            },
        ]
    )
    market_data = pd.DataFrame(columns=["ticker", "fiscal_year", "share_price", "market_cap", "eps", "bps", "per", "pbr"])
    manual_kpis = pd.DataFrame(columns=["ticker", "fiscal_year", "variable_cost_ratio", "fixed_cost_estimate"])
    return Dataset(company_master=company_master, financials=financials, market_data=market_data, manual_kpis=manual_kpis)


def test_manual_custom_report_generates_with_reference_warnings_and_diagnostic_charts():
    output_dir = PROJECT_ROOT / "work" / "pytest_reports" / uuid4().hex
    package = build_report_package(
        selected_tickers=["1001", "1002"],
        preset={
            "preset_id": "manual_custom_test",
            "name": "手動比較テスト",
            "description": "手動比較のテスト",
            "companies": ["1001", "1002"],
            "industry_mode": "strict_jpx_industry",
            "comparison_theme": "自動車・機械",
        },
        app_mode="manual_custom",
        industry_mode="strict_jpx_industry",
        dataset=_sparse_manual_dataset(),
        output_dir=output_dir,
        as_of=date(2026, 6, 18),
    )

    document = Document(package.docx_path)
    text = _document_text(package.docx_path)

    assert document.tables
    assert len(document.inline_shapes) >= 1
    assert "manual_custom（手動比較）" in text
    assert "参考判定" in text
    assert "主要財務比較は未成立" in text
    assert "未実施／推定不可" in text
    assert "販売台数" in text
    assert any("参考警告" in warning for warning in package.warnings)
    assert any(slug.endswith("_diagnostic") and path.exists() for slug, path in package.chart_paths.items())


def test_report_marks_outlier_ratios_and_keeps_them_out_of_score_notes():
    dataset = load_sample_dataset()
    dataset.financials.loc[
        (dataset.financials["ticker"] == "3543") & (dataset.financials["fiscal_year"] == dataset.financials["fiscal_year"].max()),
        "operating_income",
    ] = 10_000_000_000
    output_dir = PROJECT_ROOT / "work" / "pytest_reports" / uuid4().hex
    preset = load_presets()["friend_cafe_theme"]
    package = build_report_package(
        selected_tickers=preset["companies"],
        preset={**preset, "preset_id": "outlier_test"},
        app_mode="assignment",
        industry_mode=preset["industry_mode"],
        dataset=dataset,
        output_dir=output_dir,
        as_of=date(2026, 6, 18),
    )

    text = _document_text(package.docx_path)
    assert "異常値候補" in text
    assert "単位・スケール確認が必要" in text
